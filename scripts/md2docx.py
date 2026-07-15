#!/usr/bin/env python3
"""Markdown → Word 转换（python-docx）"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import re, sys, os

def add_code_block(doc, text):
    """添加代码块"""
    p = doc.add_paragraph()
    p.style = doc.styles['Normal']
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8.5)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    # Light gray background via shading
    shading_elm = p._element.get_or_add_pPr()
    # Just use simple formatting

def md_to_docx(md_path, docx_path):
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = '微软雅黑'
    font.size = Pt(10.5)
    style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    while i < len(lines):
        line = lines[i].rstrip()

        # Code blocks
        if line.startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                add_code_block(doc, code_text)
                code_lines = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Skip empty lines
        if not line:
            i += 1
            continue

        # Horizontal rules
        if line == '---':
            doc.add_paragraph('─' * 60)
            i += 1
            continue

        # Headings
        if line.startswith('# '):
            h = doc.add_heading(line[2:], level=1)
            i += 1
            continue
        if line.startswith('## '):
            h = doc.add_heading(line[3:], level=2)
            i += 1
            continue
        if line.startswith('### '):
            h = doc.add_heading(line[4:], level=3)
            i += 1
            continue
        if line.startswith('#### '):
            h = doc.add_heading(line[5:], level=4)
            i += 1
            continue

        # Blockquotes
        if line.startswith('> '):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(1)
            run = p.add_run(line[2:])
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Tables
        if line.startswith('|') and '|' in line[1:]:
            # Collect all table rows
            table_rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                row_line = lines[i].strip()
                if '---' in row_line and '|' in row_line:
                    i += 1
                    continue  # Skip separator row
                cells = [c.strip() for c in row_line.split('|')[1:-1]]
                table_rows.append(cells)
                i += 1

            if table_rows:
                ncols = max(len(r) for r in table_rows)
                table = doc.add_table(rows=len(table_rows), cols=ncols)
                table.style = 'Light Grid Accent 1'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER

                for ri, row_cells in enumerate(table_rows):
                    for ci, cell_text in enumerate(row_cells):
                        if ci < ncols:
                            cell = table.rows[ri].cells[ci]
                            cell.text = cell_text
                            # Make header row bold
                            if ri == 0:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.bold = True
                                        run.font.size = Pt(9)
                            else:
                                for paragraph in cell.paragraphs:
                                    for run in paragraph.runs:
                                        run.font.size = Pt(9)

                doc.add_paragraph()  # Space after table
            continue

        # Ordered list
        m = re.match(r'^(\d+)\.\s+(.*)', line)
        if m:
            p = doc.add_paragraph(m.group(2), style='List Number')
            i += 1
            continue

        # Unordered list
        if line.startswith('- ') or line.startswith('* '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            i += 1
            continue

        # Bold/italic inline
        # Bold: **text**
        line = re.sub(r'\*\*(.+?)\*\*', r'\1', line)  # strip markers, add bold later
        # Actually, we'll handle inline formatting with a simple approach

        # Regular paragraph
        p = doc.add_paragraph()
        # Simple inline formatting
        parts = re.split(r'(\*\*.+?\*\*|`.+?`|\*.+?\*)', line)
        for part in parts:
            if part.startswith('**') and part.endswith('**'):
                run = p.add_run(part[2:-2])
                run.font.bold = True
            elif part.startswith('`') and part.endswith('`'):
                run = p.add_run(part[1:-1])
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
            elif part.startswith('*') and part.endswith('*'):
                run = p.add_run(part[1:-1])
                run.font.italic = True
            else:
                p.add_run(part)

        i += 1

    doc.save(docx_path)
    print(f"Saved: {docx_path}")

if __name__ == '__main__':
    if len(sys.argv) >= 3:
        md_to_docx(sys.argv[1], sys.argv[2])
    else:
        SRC = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs', 'F2C_技术文档.md')
        DST = os.path.join(os.path.dirname(os.path.dirname(__file__)), 'docs', 'F2C_技术文档.docx')
        md_to_docx(SRC, DST)
